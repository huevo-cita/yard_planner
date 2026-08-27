#!/usr/bin/env python3
"""Convert a markdown garden plan into a .docx with real headings, lists, tables and
embedded images — the format that survives an upload to Google Docs.

    python3 build_doc.py plan.md -o plan.docx

Handles two things a plain markdown-to-docx conversion gets wrong:

1. Images. Any <img src="..." width="N" /> tag in the markdown is embedded as an inline
   picture. src resolves relative to the markdown file. width is in pixels; it is
   converted to inches at 96 dpi, capped to the text column.
2. Lists that markdown collapses. Hand-written schedules often put "- item" lines inside
   what markdown parses as a single paragraph; those are rebuilt into real <ul>/<ol>.

Then upload:
  new doc     -> gdrive.uploadFile(localPath=..., convertToGoogleFormat=True)
  update copy -> gdrive.uploadFile(localPath=..., fileId=...)   # keeps the same link
convertToGoogleFormat and fileId cannot be combined.

Requires: markdown, python-docx
"""
import argparse
import os
import re
from html.parser import HTMLParser

import markdown as md_lib
from docx import Document
from docx.shared import Inches, Pt

MAX_WIDTH_IN = 6.5
IMG_RE = re.compile(r'<img\s+[^>]*src="([^"]+)"[^>]*?/?>', re.I)
WIDTH_RE = re.compile(r'width="(\d+)"', re.I)


def build_html(md_path, images):
    """Markdown -> HTML, with <img> tags swapped for placeholder tokens."""
    src = open(md_path).read()
    src = re.sub(r'<!--.*?-->', '', src, flags=re.S)
    base = os.path.dirname(os.path.abspath(md_path))

    def swap(m):
        tag, rel = m.group(0), m.group(1)
        path = rel if os.path.isabs(rel) else os.path.join(base, rel)
        if not os.path.exists(path):
            raise SystemExit('image not found: %s (referenced as %s)' % (path, rel))
        px = WIDTH_RE.search(tag)
        width = min(int(px.group(1)) / 96.0, MAX_WIDTH_IN) if px else MAX_WIDTH_IN
        token = 'XIMAGE%dX' % len(images)
        images[token] = (path, width)
        return token

    src = IMG_RE.sub(swap, src)
    html = md_lib.markdown(src, extensions=['tables'])
    html = html.replace('<blockquote>', '').replace('</blockquote>', '')
    html = re.sub(r'<hr ?/?>', '', html)

    def fix_para(m):
        inner = m.group(1)
        if '\n' not in inner:
            return m.group(0)
        out, lead, items, kind = [], [], [], None

        def flush():
            nonlocal items, kind
            if items:
                tag = 'ul' if kind == 'ul' else 'ol'
                out.append('<%s>%s</%s>' % (tag, ''.join('<li>%s</li>' % i for i in items), tag))
                items, kind = [], None

        for ln in inner.split('\n'):
            s = ln.strip()
            mnum = re.match(r'^(\d+)\.\s+(.*)', s)
            if s.startswith('- '):
                if kind == 'ol':
                    flush()
                kind = 'ul'
                items.append(s[2:])
            elif mnum:
                if kind == 'ul':
                    flush()
                kind = 'ol'
                items.append(mnum.group(2))
            else:
                flush()
                lead.append(s)
        flush()
        prefix = '<p>%s</p>' % ' '.join(lead) if lead else ''
        return prefix + ''.join(out)

    return re.sub(r'<p>(.*?)</p>', fix_para, html, flags=re.S)


class Conv(HTMLParser):
    def __init__(self, doc, images):
        super().__init__(convert_charrefs=True)
        self.doc = doc
        self.images = images
        self.bold = 0
        self.italic = 0
        self.para = None
        self.mode = None
        self.list_style = None
        self.table_rows = None
        self.cur_row = None
        self.cur_cell = None
        self._th = False

    def emit(self, text):
        if self.table_rows is not None and self.cur_cell is not None:
            self.cur_cell.append((text, bool(self.bold), bool(self.italic)))
        elif self.para is not None:
            r = self.para.add_run(text)
            r.bold = bool(self.bold)
            r.italic = bool(self.italic)

    def handle_starttag(self, tag, attrs):
        if tag in ('h1', 'h2', 'h3', 'h4'):
            self.para = self.doc.add_heading('', level=int(tag[1]))
            self.mode = 'h'
        elif tag == 'p':
            self.para = self.doc.add_paragraph()
            self.mode = 'p'
        elif tag == 'ul':
            self.list_style = 'List Bullet'
        elif tag == 'ol':
            self.list_style = 'List Number'
        elif tag == 'li':
            self.para = self.doc.add_paragraph(style=self.list_style)
            self.mode = 'li'
        elif tag in ('strong', 'b'):
            self.bold += 1
        elif tag in ('em', 'i'):
            self.italic += 1
        elif tag == 'table':
            self.table_rows = []
        elif tag == 'tr':
            self.cur_row = []
        elif tag in ('td', 'th'):
            self.cur_cell = []
            self._th = tag == 'th'
            if self._th:
                self.bold += 1

    def handle_endtag(self, tag):
        if tag in ('h1', 'h2', 'h3', 'h4', 'p', 'li'):
            if self.para is not None and self.para.text.strip() in self.images:
                path, width = self.images[self.para.text.strip()]
                for r in list(self.para.runs):
                    r.text = ''
                self.para.add_run().add_picture(path, width=Inches(width))
            self.para = None
            self.mode = None
        elif tag in ('strong', 'b'):
            self.bold -= 1
        elif tag in ('em', 'i'):
            self.italic -= 1
        elif tag in ('td', 'th'):
            if self._th:
                self.bold -= 1
            self.cur_row.append(self.cur_cell)
            self.cur_cell = None
        elif tag == 'tr':
            self.table_rows.append(self.cur_row)
            self.cur_row = None
        elif tag == 'table':
            rows, self.table_rows = self.table_rows, None
            if not rows:
                return
            t = self.doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
            t.style = 'Table Grid'
            for i, row in enumerate(rows):
                for j, cell in enumerate(row):
                    p = t.cell(i, j).paragraphs[0]
                    for text, bold, ital in cell:
                        r = p.add_run(text)
                        r.bold, r.italic = bold, ital
                        r.font.size = Pt(9)
            self.doc.add_paragraph()

    def handle_data(self, data):
        if self.mode or (self.table_rows is not None and self.cur_cell is not None):
            self.emit(data.replace('\n', ' '))


def main():
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument('markdown', help='source .md file')
    ap.add_argument('-o', '--output', help='output .docx (default: alongside the markdown)')
    args = ap.parse_args()

    out = args.output or os.path.splitext(os.path.abspath(args.markdown))[0] + '.docx'
    images = {}
    html = build_html(args.markdown, images)

    doc = Document()
    for s in doc.sections:          # narrower margins so wide tables fit
        s.left_margin = Inches(0.6)
        s.right_margin = Inches(0.6)
    Conv(doc, images).feed(html)
    doc.save(out)

    check = Document(out)
    print(out)
    print('tables: %d | inline images: %d | paragraphs: %d'
          % (len(check.tables), len(check.inline_shapes), len(check.paragraphs)))
    if len(check.inline_shapes) != len(images):
        print('WARNING: expected %d images, embedded %d' % (len(images), len(check.inline_shapes)))


if __name__ == '__main__':
    main()
